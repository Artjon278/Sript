"""
Assignment Auto-Completer
========================
A GUI application that auto-completes .docx assignments using GitHub Models API.
Preserves original formatting, fonts, styles, and structure.

Setup:
  1. Get a GitHub Personal Access Token from https://github.com/settings/tokens
     - Create a Fine-grained token or Classic token
     - IMPORTANT: Enable the "Models" permission (read access)
  2. Paste your token into the app or set GITHUB_TOKEN environment variable
  3. Open your .docx assignment and click "Complete Assignment"

Dependencies: pip install python-docx openai requests tkinter
"""

import os
import sys
import json
import copy
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

try:
    from openai import OpenAI
except ImportError:
    print("Installing openai...")
    os.system(f"{sys.executable} -m pip install openai")
    from openai import OpenAI


# ─── Constants ────────────────────────────────────────────────────────────────
GITHUB_MODELS_ENDPOINT = "https://models.inference.ai.azure.com"
AVAILABLE_MODELS = [
    "gpt-4o",
    "gpt-4o-mini",
    "gpt-5",
    "gpt-5-mini",
    "gpt-5-nano",
    "gpt-5-chat",
    "o3-mini",
    "o4-mini",
    "DeepSeek-R1",
    "Meta-Llama-3.1-405B-Instruct",
    "Mistral-Large-2411",
]
DEFAULT_MODEL = "gpt-4o-mini"

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".assignment_completer_config.json")


# ─── Config Management ───────────────────────────────────────────────────────
def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_config(config):
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=2)
    except Exception:
        pass


# ─── Document Processing ─────────────────────────────────────────────────────
def extract_document_content(doc_path):
    """Extract text content and formatting info from a .docx file."""
    doc = Document(doc_path)
    content_blocks = []

    for i, para in enumerate(doc.paragraphs):
        block = {
            "index": i,
            "text": para.text.strip(),
            "style_name": para.style.name if para.style else "Normal",
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": [],
        }
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            }
            block["runs"].append(run_info)
        content_blocks.append(block)

    # Also extract tables
    tables_data = []
    for t_idx, table in enumerate(doc.tables):
        table_data = {"index": t_idx, "rows": []}
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data["rows"].append(row_data)
        tables_data.append(table_data)

    return content_blocks, tables_data


def build_document_text_for_ai(content_blocks, tables_data):
    """Build a readable text representation for the AI."""
    lines = []
    for block in content_blocks:
        style = block["style_name"]
        text = block["text"]
        if text:
            if "Heading" in style:
                lines.append(f"[{style}] {text}")
            else:
                lines.append(text)
        else:
            lines.append("[EMPTY LINE]")

    if tables_data:
        lines.append("\n--- TABLES ---")
        for t in tables_data:
            lines.append(f"\nTable {t['index'] + 1}:")
            for row in t["rows"]:
                lines.append(" | ".join(row))

    return "\n".join(lines)


def complete_assignment_with_ai(token, model, doc_text, instructions, progress_callback=None):
    """Send document to AI and get completed version."""
    client = OpenAI(
        base_url=GITHUB_MODELS_ENDPOINT,
        api_key=token,
    )

    system_prompt = """You are an expert academic assignment completer. You help students complete their assignments.

CRITICAL RULES:
1. You receive a document with questions, prompts, or incomplete sections.
2. You MUST return the COMPLETE document with all answers filled in.
3. PRESERVE the EXACT same structure - headings, numbering, formatting markers.
4. Keep all original text exactly as-is, only ADD answers/completions where needed.
5. If a line has a question, write the answer on the next line(s) after it.
6. If there are blank spaces or underlines (___) meant for answers, fill them in.
7. If a table has empty cells, fill them with appropriate content.
8. Match the academic level and tone of the existing content.
9. Write answers in the SAME LANGUAGE as the questions (Albanian or English).
10. Be thorough but concise - match what a good student would write.
11. For each section/paragraph, output it in this exact format:
    [PARA:index_number] content here
    Where index_number matches the original paragraph index.
12. For NEW answer paragraphs that you add, use:
    [NEW_AFTER:index_number] answer content here
    This means insert this new paragraph after the paragraph with that index.
13. For tables, use:
    [TABLE:table_index] 
    row1_col1 | row1_col2 | ...
    row2_col1 | row2_col2 | ...

IMPORTANT: Return ALL paragraphs (including unchanged ones) so the document stays complete."""

    user_message = f"""Here is my assignment document. Please complete it following all the rules.

{f"Additional instructions: {instructions}" if instructions else ""}

--- DOCUMENT START ---
{doc_text}
--- DOCUMENT END ---

Please complete this assignment. Return the full document with answers filled in, using the [PARA:N] format for each paragraph."""

    if progress_callback:
        progress_callback("Sending to AI model...")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        temperature=0.3,
        max_tokens=16000,
    )

    return response.choices[0].message.content


def parse_ai_response(ai_text):
    """Parse the AI response into structured paragraphs."""
    paragraphs = {}
    new_paragraphs = {}  # {after_index: [list of new paragraph texts]}
    tables = {}

    current_type = None
    current_index = None
    current_lines = []

    for line in ai_text.split("\n"):
        stripped = line.strip()

        # Check for paragraph markers
        if stripped.startswith("[PARA:"):
            # Save previous
            if current_type == "para" and current_index is not None:
                paragraphs[current_index] = "\n".join(current_lines).strip()
            elif current_type == "new" and current_index is not None:
                if current_index not in new_paragraphs:
                    new_paragraphs[current_index] = []
                new_paragraphs[current_index].append("\n".join(current_lines).strip())

            try:
                idx_str = stripped.split("]")[0].replace("[PARA:", "")
                current_index = int(idx_str)
                current_type = "para"
                content_after = stripped.split("]", 1)[1].strip() if "]" in stripped else ""
                current_lines = [content_after] if content_after else []
            except (ValueError, IndexError):
                current_lines.append(line)

        elif stripped.startswith("[NEW_AFTER:"):
            # Save previous
            if current_type == "para" and current_index is not None:
                paragraphs[current_index] = "\n".join(current_lines).strip()
            elif current_type == "new" and current_index is not None:
                if current_index not in new_paragraphs:
                    new_paragraphs[current_index] = []
                new_paragraphs[current_index].append("\n".join(current_lines).strip())

            try:
                idx_str = stripped.split("]")[0].replace("[NEW_AFTER:", "")
                current_index = int(idx_str)
                current_type = "new"
                content_after = stripped.split("]", 1)[1].strip() if "]" in stripped else ""
                current_lines = [content_after] if content_after else []
            except (ValueError, IndexError):
                current_lines.append(line)

        elif stripped.startswith("[TABLE:"):
            # Save previous
            if current_type == "para" and current_index is not None:
                paragraphs[current_index] = "\n".join(current_lines).strip()
            elif current_type == "new" and current_index is not None:
                if current_index not in new_paragraphs:
                    new_paragraphs[current_index] = []
                new_paragraphs[current_index].append("\n".join(current_lines).strip())

            try:
                idx_str = stripped.split("]")[0].replace("[TABLE:", "")
                current_index = int(idx_str)
                current_type = "table"
                current_lines = []
            except (ValueError, IndexError):
                current_lines.append(line)
        else:
            current_lines.append(line)

    # Save last block
    if current_type == "para" and current_index is not None:
        paragraphs[current_index] = "\n".join(current_lines).strip()
    elif current_type == "new" and current_index is not None:
        if current_index not in new_paragraphs:
            new_paragraphs[current_index] = []
        new_paragraphs[current_index].append("\n".join(current_lines).strip())

    if current_type == "table" and current_index is not None:
        table_rows = []
        for l in current_lines:
            l = l.strip()
            if l and "|" in l:
                table_rows.append([c.strip() for c in l.split("|")])
        tables[current_index] = table_rows

    return paragraphs, new_paragraphs, tables


def apply_completion_to_document(doc_path, output_path, ai_response_text, progress_callback=None):
    """Apply AI completions back to the document, preserving formatting."""
    doc = Document(doc_path)

    if progress_callback:
        progress_callback("Parsing AI response...")

    parsed_paras, new_paras, parsed_tables = parse_ai_response(ai_response_text)

    # If parsing found structured data, apply it
    if parsed_paras or new_paras:
        if progress_callback:
            progress_callback("Applying completions to document...")

        # Update existing paragraphs
        for idx, new_text in parsed_paras.items():
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                old_text = para.text.strip()
                if old_text != new_text and new_text:
                    # Preserve formatting of the first run, update text
                    if para.runs:
                        # Keep formatting from original runs
                        first_run = para.runs[0]
                        # Clear all runs
                        for run in para.runs[1:]:
                            run.text = ""
                        first_run.text = new_text
                    else:
                        para.text = new_text

        # Insert new paragraphs (process in reverse order to maintain indices)
        sorted_new = sorted(new_paras.items(), key=lambda x: x[0], reverse=True)
        for after_idx, texts in sorted_new:
            if after_idx < len(doc.paragraphs):
                ref_para = doc.paragraphs[after_idx]
                # Get formatting from reference paragraph or nearby
                ref_style = ref_para.style
                ref_run_format = None
                if ref_para.runs:
                    ref_run_format = ref_para.runs[0]

                for text in reversed(texts):
                    if not text:
                        continue
                    # Insert paragraph after the reference
                    new_para_element = copy.deepcopy(ref_para._element)
                    ref_para._element.addnext(new_para_element)

                    # Find the new paragraph in doc and set text
                    # We need to get the paragraph object for the new element
                    new_para = None
                    for p in doc.paragraphs:
                        if p._element is new_para_element:
                            new_para = p
                            break

                    if new_para:
                        # Use Normal style for answers (not heading style)
                        if ref_style and "Heading" in ref_style.name:
                            try:
                                new_para.style = doc.styles["Normal"]
                            except Exception:
                                pass

                        if new_para.runs:
                            new_para.runs[0].text = text
                            for run in new_para.runs[1:]:
                                run.text = ""
                            # Remove bold from answer paragraphs
                            for run in new_para.runs:
                                run.bold = False
                        else:
                            new_para.text = text

        # Update tables
        for t_idx, rows_data in parsed_tables.items():
            if t_idx < len(doc.tables):
                table = doc.tables[t_idx]
                for r_idx, row_data in enumerate(rows_data):
                    if r_idx < len(table.rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(table.rows[r_idx].cells):
                                cell = table.rows[r_idx].cells[c_idx]
                                if not cell.text.strip() and cell_text.strip():
                                    cell.text = cell_text

    else:
        # Fallback: Simple approach - use AI response as a guide
        # Try a simpler completion strategy
        if progress_callback:
            progress_callback("Using fallback completion strategy...")

        _apply_simple_completion(doc, ai_response_text)

    if progress_callback:
        progress_callback("Saving document...")

    doc.save(output_path)
    return output_path


def _apply_simple_completion(doc, ai_text):
    """Fallback: Insert AI-generated answers after empty/question paragraphs."""
    ai_lines = [l.strip() for l in ai_text.split("\n") if l.strip()]

    # Find paragraphs that look like questions or have blank answers
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Check if this is a question without an answer following it
        is_question = (
            text.endswith("?") or
            text.endswith(":") or
            "___" in text or
            "..." in text
        )

        if is_question and "___" in text:
            # Try to find the answer in AI text that corresponds
            for ai_line in ai_lines:
                if any(word in ai_line.lower() for word in text.lower().split()[:3]):
                    # Replace blanks with answer
                    if para.runs:
                        full_text = para.text.replace("___", ai_line.split(":")[-1].strip() if ":" in ai_line else ai_line)
                        para.runs[0].text = full_text
                        for run in para.runs[1:]:
                            run.text = ""
                    break


# ─── GUI Application ─────────────────────────────────────────────────────────
class AssignmentCompleterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("📝 Assignment Auto-Completer")
        self.root.geometry("900x720")
        self.root.minsize(750, 600)

        # Configure theme colors
        self.bg_color = "#1e1e2e"
        self.fg_color = "#cdd6f4"
        self.accent_color = "#89b4fa"
        self.accent2_color = "#a6e3a1"
        self.surface_color = "#313244"
        self.surface2_color = "#45475a"
        self.red_color = "#f38ba8"
        self.yellow_color = "#f9e2af"

        self.root.configure(bg=self.bg_color)

        # Style
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure("TFrame", background=self.bg_color)
        self.style.configure("TLabel", background=self.bg_color, foreground=self.fg_color, font=("Segoe UI", 10))
        self.style.configure("Title.TLabel", font=("Segoe UI", 18, "bold"), foreground=self.accent_color)
        self.style.configure("Subtitle.TLabel", font=("Segoe UI", 10), foreground=self.surface2_color)
        self.style.configure("TButton", font=("Segoe UI", 10), padding=8)
        self.style.configure("Accent.TButton", font=("Segoe UI", 11, "bold"), padding=12)
        self.style.configure("TCombobox", font=("Segoe UI", 10))
        self.style.configure("TLabelframe", background=self.bg_color, foreground=self.fg_color)
        self.style.configure("TLabelframe.Label", background=self.bg_color, foreground=self.accent_color, font=("Segoe UI", 10, "bold"))

        self.config = load_config()
        self.file_path = None
        self.is_processing = False

        self._build_ui()

    def _build_ui(self):
        # Main container
        main = ttk.Frame(self.root, padding=20)
        main.pack(fill=tk.BOTH, expand=True)

        # ── Header ──
        header = ttk.Frame(main)
        header.pack(fill=tk.X, pady=(0, 15))
        ttk.Label(header, text="📝 Assignment Auto-Completer", style="Title.TLabel").pack(side=tk.LEFT)
        ttk.Label(header, text="Powered by GitHub Models", style="Subtitle.TLabel").pack(side=tk.RIGHT, pady=(8, 0))

        # ── API Settings Frame ──
        api_frame = ttk.LabelFrame(main, text="⚙️  API Settings", padding=12)
        api_frame.pack(fill=tk.X, pady=(0, 10))

        # Token row
        token_row = ttk.Frame(api_frame)
        token_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Label(token_row, text="GitHub Token:", width=14).pack(side=tk.LEFT)
        self.token_var = tk.StringVar(value=self.config.get("token", os.environ.get("GITHUB_TOKEN", "")))
        self.token_entry = tk.Entry(
            token_row, textvariable=self.token_var, show="●",
            font=("Consolas", 10), bg=self.surface_color, fg=self.fg_color,
            insertbackground=self.fg_color, relief=tk.FLAT, bd=5
        )
        self.token_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))

        self.show_token_var = tk.BooleanVar(value=False)
        tk.Checkbutton(
            token_row, text="Show", variable=self.show_token_var,
            command=self._toggle_token_visibility,
            bg=self.bg_color, fg=self.fg_color, selectcolor=self.surface_color,
            activebackground=self.bg_color, activeforeground=self.fg_color
        ).pack(side=tk.RIGHT)

        # Model row
        model_row = ttk.Frame(api_frame)
        model_row.pack(fill=tk.X)
        ttk.Label(model_row, text="AI Model:", width=14).pack(side=tk.LEFT)
        self.model_var = tk.StringVar(value=self.config.get("model", DEFAULT_MODEL))
        model_combo = ttk.Combobox(
            model_row, textvariable=self.model_var, values=AVAILABLE_MODELS,
            width=35
        )
        model_combo.pack(side=tk.LEFT, padx=(5, 0))

        save_btn = tk.Button(
            model_row, text="💾 Save Settings", command=self._save_settings,
            bg=self.surface2_color, fg=self.fg_color, relief=tk.FLAT,
            font=("Segoe UI", 9), padx=10, pady=3, cursor="hand2"
        )
        save_btn.pack(side=tk.RIGHT)

        # ── Document Frame ──
        doc_frame = ttk.LabelFrame(main, text="📄  Document", padding=12)
        doc_frame.pack(fill=tk.X, pady=(0, 10))

        file_row = ttk.Frame(doc_frame)
        file_row.pack(fill=tk.X, pady=(0, 8))

        self.file_label = tk.Label(
            file_row, text="No file selected", anchor=tk.W,
            bg=self.surface_color, fg=self.surface2_color,
            font=("Consolas", 10), relief=tk.FLAT, bd=5, padx=10, pady=5
        )
        self.file_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

        browse_btn = tk.Button(
            file_row, text="📂 Browse", command=self._browse_file,
            bg=self.accent_color, fg="#1e1e2e", relief=tk.FLAT,
            font=("Segoe UI", 10, "bold"), padx=15, pady=5, cursor="hand2"
        )
        browse_btn.pack(side=tk.RIGHT, padx=(8, 0))

        # Output option
        output_row = ttk.Frame(doc_frame)
        output_row.pack(fill=tk.X)
        self.overwrite_var = tk.BooleanVar(value=False)
        tk.Checkbutton(
            output_row, text="Overwrite original file (otherwise saves as *_completed.docx)",
            variable=self.overwrite_var,
            bg=self.bg_color, fg=self.fg_color, selectcolor=self.surface_color,
            activebackground=self.bg_color, activeforeground=self.fg_color,
            font=("Segoe UI", 9)
        ).pack(side=tk.LEFT)

        # ── Instructions Frame ──
        instr_frame = ttk.LabelFrame(main, text="📋  Extra Instructions (optional)", padding=12)
        instr_frame.pack(fill=tk.X, pady=(0, 10))

        self.instructions_text = tk.Text(
            instr_frame, height=3, wrap=tk.WORD,
            bg=self.surface_color, fg=self.fg_color,
            insertbackground=self.fg_color, relief=tk.FLAT, bd=5,
            font=("Segoe UI", 10)
        )
        self.instructions_text.pack(fill=tk.X)
        self.instructions_text.insert("1.0", 'e.g. "Answer in Albanian", "Keep answers short", "Use formal language"')
        self.instructions_text.bind("<FocusIn>", self._clear_placeholder)
        self._instructions_placeholder = True

        # ── Action Button ──
        btn_frame = ttk.Frame(main)
        btn_frame.pack(fill=tk.X, pady=(0, 10))

        self.complete_btn = tk.Button(
            btn_frame, text="🚀  Complete Assignment", command=self._start_completion,
            bg=self.accent2_color, fg="#1e1e2e", relief=tk.FLAT,
            font=("Segoe UI", 13, "bold"), padx=20, pady=12, cursor="hand2",
            activebackground="#74c790"
        )
        self.complete_btn.pack(fill=tk.X)

        # ── Progress / Log ──
        log_frame = ttk.LabelFrame(main, text="📊  Progress", padding=12)
        log_frame.pack(fill=tk.BOTH, expand=True)

        self.progress_bar = ttk.Progressbar(log_frame, mode="indeterminate", length=300)
        self.progress_bar.pack(fill=tk.X, pady=(0, 8))

        self.log_text = scrolledtext.ScrolledText(
            log_frame, height=8, wrap=tk.WORD, state=tk.DISABLED,
            bg=self.surface_color, fg=self.fg_color,
            insertbackground=self.fg_color, relief=tk.FLAT, bd=5,
            font=("Consolas", 9)
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)

    def _toggle_token_visibility(self):
        self.token_entry.config(show="" if self.show_token_var.get() else "●")

    def _clear_placeholder(self, event=None):
        if self._instructions_placeholder:
            self.instructions_text.delete("1.0", tk.END)
            self._instructions_placeholder = False

    def _save_settings(self):
        self.config["token"] = self.token_var.get().strip()
        self.config["model"] = self.model_var.get()
        save_config(self.config)
        self._log("✅ Settings saved!")

    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Select Assignment Document",
            filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")]
        )
        if path:
            self.file_path = path
            display_name = os.path.basename(path)
            self.file_label.config(text=f"  {display_name}", fg=self.accent_color)
            self._log(f"📄 Selected: {display_name}")

    def _log(self, message):
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.log_text.config(state=tk.DISABLED)

    def _set_processing(self, state):
        self.is_processing = state
        if state:
            self.complete_btn.config(state=tk.DISABLED, text="⏳  Processing...", bg=self.surface2_color)
            self.progress_bar.start(15)
        else:
            self.complete_btn.config(state=tk.NORMAL, text="🚀  Complete Assignment", bg=self.accent2_color)
            self.progress_bar.stop()

    def _start_completion(self):
        # Validate inputs
        token = self.token_var.get().strip()
        if not token:
            messagebox.showerror("Error", "Please enter your GitHub Personal Access Token.\n\nGet one at: https://github.com/settings/tokens")
            return
        if not self.file_path:
            messagebox.showerror("Error", "Please select a .docx file first.")
            return
        if not os.path.exists(self.file_path):
            messagebox.showerror("Error", "Selected file not found.")
            return

        self._set_processing(True)
        thread = threading.Thread(target=self._run_completion, daemon=True)
        thread.start()

    def _run_completion(self):
        try:
            token = self.token_var.get().strip()
            model = self.model_var.get()
            instructions = self.instructions_text.get("1.0", tk.END).strip()
            if self._instructions_placeholder:
                instructions = ""

            # Step 1: Extract document content
            self.root.after(0, self._log, "📖 Reading document...")
            content_blocks, tables_data = extract_document_content(self.file_path)
            self.root.after(0, self._log, f"   Found {len(content_blocks)} paragraphs, {len(tables_data)} tables")

            # Step 2: Build text for AI
            doc_text = build_document_text_for_ai(content_blocks, tables_data)
            self.root.after(0, self._log, f"📝 Document text prepared ({len(doc_text)} chars)")

            # Step 3: Send to AI
            self.root.after(0, self._log, f"🤖 Sending to {model}...")

            def progress_cb(msg):
                self.root.after(0, self._log, f"   {msg}")

            ai_response = complete_assignment_with_ai(token, model, doc_text, instructions, progress_cb)
            self.root.after(0, self._log, f"✅ AI response received ({len(ai_response)} chars)")

            # Step 4: Apply to document
            if self.overwrite_var.get():
                output_path = self.file_path
            else:
                base, ext = os.path.splitext(self.file_path)
                output_path = f"{base}_completed{ext}"

            self.root.after(0, self._log, "📝 Applying completions to document...")
            result_path = apply_completion_to_document(self.file_path, output_path, ai_response, progress_cb)

            self.root.after(0, self._log, f"🎉 Done! Saved to: {os.path.basename(result_path)}")
            self.root.after(0, self._log, f"   Full path: {result_path}")

            # Ask to open
            self.root.after(0, lambda: self._ask_open_file(result_path))

        except Exception as e:
            error_msg = str(e)
            if "401" in error_msg or "unauthorized" in error_msg.lower():
                if "models" in error_msg.lower() and "permission" in error_msg.lower():
                    self.root.after(0, self._log, "❌ Your token is missing the 'Models' permission!")
                    self.root.after(0, self._log, "   Go to: https://github.com/settings/tokens")
                    self.root.after(0, self._log, "   Edit your token → enable 'Models' (read) permission → Save")
                else:
                    self.root.after(0, self._log, "❌ Invalid GitHub token. Please check your token and try again.")
            elif "404" in error_msg or "unknown_model" in error_msg.lower() or "Unknown model" in error_msg:
                self.root.after(0, self._log, f"❌ Model '{self.model_var.get()}' does not exist on GitHub Models!")
                self.root.after(0, self._log, f"   Try one of: gpt-4o, gpt-4o-mini, gpt-5, gpt-5-mini, o4-mini")
            elif "400" in error_msg:
                self.root.after(0, self._log, f"❌ Bad request: {error_msg}")
            elif "rate" in error_msg.lower():
                self.root.after(0, self._log, "❌ Rate limit reached. Please wait a moment and try again.")
            else:
                self.root.after(0, self._log, f"❌ Error: {error_msg}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Completion failed:\n{error_msg}"))
        finally:
            self.root.after(0, self._set_processing, False)

    def _ask_open_file(self, path):
        if messagebox.askyesno("Success!", f"Assignment completed!\n\nSaved to:\n{os.path.basename(path)}\n\nOpen the file now?"):
            os.startfile(path)


# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    root = tk.Tk()

    # Set window icon (optional)
    try:
        root.iconbitmap(default="")
    except Exception:
        pass

    # Center window on screen
    root.update_idletasks()
    w, h = 900, 720
    x = (root.winfo_screenwidth() // 2) - (w // 2)
    y = (root.winfo_screenheight() // 2) - (h // 2)
    root.geometry(f"{w}x{h}+{x}+{y}")

    app = AssignmentCompleterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
